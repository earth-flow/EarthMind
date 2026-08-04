"""Sandboxed Word (.docx) document tool component exposing structured edit tools to agents.

Composes ``FileSystemToolComponent`` rather than duplicating its sandbox
enforcement: path validation (traversal/reserved-segment/deny-list/hardlink
checks) and the TOCTOU-safe, symlink-refusing byte read/write primitives are
security-critical and already tested there. This module only adds
docx-specific structure (paragraphs, styles, tables) on top of raw bytes.
"""

from __future__ import annotations

import io
import json
import zipfile
from typing import TYPE_CHECKING

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from langchain_core.tools import StructuredTool
from pydantic import BaseModel, Field

from lfx.components.files_and_knowledge.filesystem import (
    MAX_FILE_SIZE_BYTES,
    FileSystemToolComponent,
    _read_bytes_no_follow,
    _write_bytes_no_follow,
)
from lfx.custom.custom_component.component import Component
from lfx.inputs.inputs import BoolInput, StrInput
from lfx.io import Output
from lfx.schema.data import Data

if TYPE_CHECKING:
    from pathlib import Path

    from lfx.field_typing import Tool

# Paragraph styles the agent may request. Restricting to a known-good set
# (rather than passing any string straight to python-docx) means a bad style
# name fails with a structured, comprehensible error instead of a raw
# python-docx KeyError, and keeps output documents visually consistent
# regardless of which template/theme the .docx was created from.
_ALLOWED_STYLES = frozenset(
    {
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "List Bullet",
        "List Number",
        "Quote",
        "Caption",
    }
)

MAX_TABLE_CELLS = 2000  # generous ceiling against an agent requesting a pathological rows*cols grid


class _DocxPathArgs(BaseModel):
    path: str = Field(..., description="Path to the .docx file, relative to the sandbox root.")


class _DocxCreateArgs(BaseModel):
    path: str = Field(..., description="Path to the new .docx file, relative to the sandbox root.")
    title: str | None = Field(default=None, description="Optional title, inserted as a 'Title'-styled paragraph.")
    overwrite: bool = Field(default=False, description="If false (default), fails when the file already exists.")


class _DocxAppendParagraphArgs(BaseModel):
    path: str = Field(..., description="Path to the .docx file, relative to the sandbox root.")
    text: str = Field(..., description="Paragraph text to append at the end of the document.")
    style: str | None = Field(
        default=None,
        description=f"Optional paragraph style. One of: {sorted(_ALLOWED_STYLES)}.",
    )


class _DocxReplaceTextArgs(BaseModel):
    path: str = Field(..., description="Path to the .docx file, relative to the sandbox root.")
    old_text: str = Field(..., description="Exact text to find (matched against each paragraph's full text).")
    new_text: str = Field(..., description="Replacement text for the whole matching paragraph.")
    replace_all: bool = Field(
        default=False,
        description="Replace every matching paragraph instead of failing when more than one matches.",
    )


class _DocxAddTableArgs(BaseModel):
    path: str = Field(..., description="Path to the .docx file, relative to the sandbox root.")
    rows: int = Field(..., description="Number of rows.")
    cols: int = Field(..., description="Number of columns.")
    data: list[list[str]] | None = Field(
        default=None,
        description="Optional cell contents as a list of rows, each a list of strings. Missing cells stay empty.",
    )
    style: str | None = Field(default=None, description="Optional Word table style name, e.g. 'Light Grid'.")


class WordDocumentToolComponent(Component):
    display_name = "Word Document"
    description = "Sandboxed Word (.docx) document read/edit tools for agents."
    icon = "file-text"
    name = "WordDocumentTool"

    add_tool_output = True

    inputs = [
        StrInput(
            name="root_path",
            display_name="Workspace Sub-path",
            required=False,
            value="",
            info=(
                "Sub-folder inside your sandboxed workspace. Leave empty to use the root "
                "of the workspace (the same sandbox the File System tool operates on)."
            ),
        ),
        BoolInput(
            name="read_only",
            display_name="Read Only",
            value=False,
            advanced=True,
            info="If true, create/append/replace/table operations are disabled.",
        ),
        # See filesystem.py's identical field for why this hidden input exists:
        # it is the only way to surface the "Tool Mode" toggle on the node header.
        StrInput(
            name="tool_mode_trigger",
            display_name="",
            show=False,
            tool_mode=True,
            required=False,
        ),
    ]

    outputs = [
        Output(display_name="JSON", name="metadata", method="build_metadata", types=["Data"]),
    ]

    def build_metadata(self) -> Data:
        registered = ["docx_read"]
        if not self.read_only:
            registered.extend(["docx_create", "docx_append_paragraph", "docx_replace_text", "docx_add_table"])
        return Data(
            data={"root_path": self.root_path, "read_only": bool(self.read_only), "tools_registered": registered}
        )

    def _fs(self) -> FileSystemToolComponent:
        """Build a FileSystemToolComponent bound to the same identity/scope as this component.

        Only used for its (well-tested) ``_validate_root``/``_validate_path``
        boundary checks — never for its own tool functions. ``build_toolkit()``
        sets ``root_path``/``_user_id``/``_force_isolation`` on this component the
        same way it does on the real FileSystemToolComponent instance, so both
        tools resolve to the identical per-user sandbox root.
        """
        fs = FileSystemToolComponent(root_path=self.root_path, read_only=self.read_only)
        fs._user_id = getattr(self, "_user_id", None)  # noqa: SLF001 — intentional cross-component sandbox reuse
        fs._force_isolation = getattr(self, "_force_isolation", False)  # noqa: SLF001
        return fs

    async def _get_tools(self) -> list[Tool]:
        tools: list[Tool] = [self._make_read_tool()]
        if not self.read_only:
            tools.append(self._make_create_tool())
            tools.append(self._make_append_paragraph_tool())
            tools.append(self._make_replace_text_tool())
            tools.append(self._make_add_table_tool())
        return tools

    def _validate_docx_path(self, path: str) -> Path | dict:
        """Resolve + validate `path`, additionally requiring a `.docx` extension.

        Returns the resolved Path on success, or a structured error dict.
        """
        if not path.lower().endswith(".docx"):
            return {"error": f"Path must end with .docx: {path}", "path": path}
        try:
            return self._fs()._validate_path(path)  # noqa: SLF001 — intentional cross-component sandbox reuse
        except PermissionError as exc:
            return {"error": str(exc), "path": path}

    def _load_document(self, resolved: Path, path: str) -> Document | dict:
        if not resolved.exists():
            return {"error": f"File not found: {path}", "path": path}
        try:
            size = resolved.stat().st_size
        except OSError as exc:
            return {"error": f"Cannot stat file: {exc.strerror or exc}", "path": path}
        if size > MAX_FILE_SIZE_BYTES:
            return {"error": f"File size {size} exceeds limit of {MAX_FILE_SIZE_BYTES} bytes", "path": path}
        try:
            raw = _read_bytes_no_follow(resolved)
        except OSError as exc:
            return {"error": f"Cannot read file: {exc.strerror or exc}", "path": path}
        try:
            return Document(io.BytesIO(raw))
        except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError) as exc:
            return {"error": f"Not a valid .docx file: {exc}", "path": path}

    def _save_document(self, document: Document, resolved: Path, path: str) -> dict:
        buffer = io.BytesIO()
        document.save(buffer)
        encoded = buffer.getvalue()
        if len(encoded) > MAX_FILE_SIZE_BYTES:
            return {
                "error": f"Resulting file size {len(encoded)} exceeds limit of {MAX_FILE_SIZE_BYTES} bytes",
                "path": path,
            }
        try:
            resolved.parent.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            return {"error": f"Cannot create parent directory: {exc.strerror or exc}", "path": path}
        try:
            _write_bytes_no_follow(resolved, encoded)
        except OSError as exc:
            return {"error": f"Cannot write file: {exc.strerror or exc}", "path": path}
        return {"status": "ok", "path": path, "bytes_written": len(encoded)}

    # -- docx_read ---------------------------------------------------------

    def _docx_read(self, path: str) -> dict:
        resolved = self._validate_docx_path(path)
        if isinstance(resolved, dict):
            return resolved
        document = self._load_document(resolved, path)
        if isinstance(document, dict):
            return document

        paragraphs = [
            {"index": i, "style": p.style.name if p.style else None, "text": p.text}
            for i, p in enumerate(document.paragraphs)
            if p.text.strip()
        ]
        tables = [
            {"index": i, "rows": [[cell.text for cell in row.cells] for row in table.rows]}
            for i, table in enumerate(document.tables)
        ]
        return {"status": "ok", "path": path, "paragraphs": paragraphs, "tables": tables}

    def _make_read_tool(self) -> StructuredTool:
        def _run(path: str) -> str:
            return json.dumps(self._docx_read(path))

        return StructuredTool.from_function(
            name="docx_read",
            description=(
                "Read a Word (.docx) file's structure: non-empty paragraphs (with style and text) "
                "and tables (as rows of cell text)."
            ),
            func=_run,
            args_schema=_DocxPathArgs,
            tags=["docx_read"],
        )

    # -- docx_create ---------------------------------------------------------

    def _docx_create(self, path: str, title: str | None = None, *, overwrite: bool = False) -> dict:
        resolved = self._validate_docx_path(path)
        if isinstance(resolved, dict):
            return resolved
        if resolved.exists() and not overwrite:
            return {"error": f"File already exists: {path}. Pass overwrite=True to replace it.", "path": path}

        document = Document()
        if title:
            document.add_heading(title, level=0)
        return self._save_document(document, resolved, path)

    def _make_create_tool(self) -> StructuredTool:
        def _run(path: str, title: str | None = None, *, overwrite: bool = False) -> str:
            return json.dumps(self._docx_create(path, title=title, overwrite=overwrite))

        return StructuredTool.from_function(
            name="docx_create",
            description="Create a new Word (.docx) file, optionally with a title.",
            func=_run,
            args_schema=_DocxCreateArgs,
            tags=["docx_create"],
        )

    # -- docx_append_paragraph -----------------------------------------------

    def _docx_append_paragraph(self, path: str, text: str, style: str | None = None) -> dict:
        if style is not None and style not in _ALLOWED_STYLES:
            return {"error": f"Unknown style {style!r}. Must be one of {sorted(_ALLOWED_STYLES)}", "path": path}
        resolved = self._validate_docx_path(path)
        if isinstance(resolved, dict):
            return resolved
        document = self._load_document(resolved, path)
        if isinstance(document, dict):
            return document
        document.add_paragraph(text, style=style)
        return self._save_document(document, resolved, path)

    def _make_append_paragraph_tool(self) -> StructuredTool:
        def _run(path: str, text: str, style: str | None = None) -> str:
            return json.dumps(self._docx_append_paragraph(path, text, style=style))

        return StructuredTool.from_function(
            name="docx_append_paragraph",
            description="Append a paragraph to the end of an existing Word (.docx) file.",
            func=_run,
            args_schema=_DocxAppendParagraphArgs,
            tags=["docx_append_paragraph"],
        )

    # -- docx_replace_text ----------------------------------------------------

    def _docx_replace_text(self, path: str, old_text: str, new_text: str, *, replace_all: bool = False) -> dict:
        """Replace whole paragraphs whose text contains `old_text`.

        Rewrites the matching paragraph's runs into a single run carrying the
        substituted text. This deliberately trades run-level formatting
        granularity (bold/italic spans mid-paragraph) for a replacement that
        is always well-defined, regardless of how python-docx split the
        original text into runs — matching the same "exact match, ambiguity
        fails closed unless replace_all" contract as the plain-text edit_file
        tool for consistency.
        """
        if old_text == "":
            return {"error": "old_text must not be empty", "path": path}
        resolved = self._validate_docx_path(path)
        if isinstance(resolved, dict):
            return resolved
        document = self._load_document(resolved, path)
        if isinstance(document, dict):
            return document

        matching = [p for p in document.paragraphs if old_text in p.text]
        if not matching:
            return {"error": f"old_text not found in any paragraph: {path}", "path": path}
        if len(matching) > 1 and not replace_all:
            return {
                "error": (
                    f"old_text is ambiguous: matches {len(matching)} paragraphs in {path}. "
                    "Pass replace_all=True to replace every match."
                ),
                "path": path,
                "matches": len(matching),
            }

        for paragraph in matching:
            new_paragraph_text = paragraph.text.replace(old_text, new_text)
            for run in list(paragraph.runs):
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_paragraph_text
            else:
                paragraph.add_run(new_paragraph_text)

        result = self._save_document(document, resolved, path)
        if "error" not in result:
            result["replacements"] = len(matching)
        return result

    def _make_replace_text_tool(self) -> StructuredTool:
        def _run(path: str, old_text: str, new_text: str, *, replace_all: bool = False) -> str:
            return json.dumps(self._docx_replace_text(path, old_text, new_text, replace_all=replace_all))

        return StructuredTool.from_function(
            name="docx_replace_text",
            description=(
                "Replace text within a Word (.docx) file's paragraphs. Matches whole paragraphs "
                "containing old_text and rewrites them with old_text substituted for new_text "
                "(the matched paragraph's run-level formatting is flattened into a single run). "
                "Fails on ambiguous matches (multiple paragraphs) unless replace_all=True."
            ),
            func=_run,
            args_schema=_DocxReplaceTextArgs,
            tags=["docx_replace_text"],
        )

    # -- docx_add_table -------------------------------------------------------

    def _docx_add_table(
        self, path: str, rows: int, cols: int, data: list[list[str]] | None = None, style: str | None = None
    ) -> dict:
        if rows <= 0 or cols <= 0:
            return {"error": "rows and cols must both be positive integers", "path": path}
        if rows * cols > MAX_TABLE_CELLS:
            return {"error": f"Table of {rows}x{cols} exceeds the {MAX_TABLE_CELLS}-cell limit", "path": path}
        resolved = self._validate_docx_path(path)
        if isinstance(resolved, dict):
            return resolved
        document = self._load_document(resolved, path)
        if isinstance(document, dict):
            return document

        try:
            table = document.add_table(rows=rows, cols=cols)
            if style:
                table.style = style
        except KeyError as exc:
            return {"error": f"Unknown table style {style!r}: {exc}", "path": path}

        if data:
            for r, row_values in enumerate(data[:rows]):
                for c, value in enumerate(row_values[:cols]):
                    table.rows[r].cells[c].text = str(value)

        return self._save_document(document, resolved, path)

    def _make_add_table_tool(self) -> StructuredTool:
        def _run(path: str, rows: int, cols: int, data: list[list[str]] | None = None, style: str | None = None) -> str:
            return json.dumps(self._docx_add_table(path, rows, cols, data=data, style=style))

        return StructuredTool.from_function(
            name="docx_add_table",
            description="Append a table to the end of an existing Word (.docx) file, optionally pre-filled with data.",
            func=_run,
            args_schema=_DocxAddTableArgs,
            tags=["docx_add_table"],
        )
