from __future__ import annotations

import asyncio
import io
import json
import os
import re
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import httpx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from lfx.base.data.utils import extract_text_from_bytes
from lfx.log import logger

PARSER_STRATEGY_AUTO = "auto"
PARSER_STRATEGY_MINERU_MARKDOWN = "mineru_markdown"
PARSER_STRATEGY_PLAIN_TEXT = "plain_text"

CHUNK_STRATEGY_AUTO = "auto"
CHUNK_STRATEGY_HEADING_MARKDOWN = "heading_markdown"
CHUNK_STRATEGY_RECURSIVE_TEXT = "recursive_text"

SUPPORTED_PARSER_STRATEGIES = frozenset(
    {
        PARSER_STRATEGY_AUTO,
        PARSER_STRATEGY_MINERU_MARKDOWN,
        PARSER_STRATEGY_PLAIN_TEXT,
    }
)
SUPPORTED_CHUNK_STRATEGIES = frozenset(
    {
        CHUNK_STRATEGY_AUTO,
        CHUNK_STRATEGY_HEADING_MARKDOWN,
        CHUNK_STRATEGY_RECURSIVE_TEXT,
    }
)

_SOFFICE_CANDIDATES = (
    "libreoffice",
    "soffice",
    "/usr/bin/libreoffice",
    "/usr/bin/soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    r"C:\Program Files\LibreOffice\program\soffice.exe",
)
_HEADING_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$")


@dataclass(slots=True)
class ProcessedChunk:
    content: str
    index: int
    title: str = ""
    level: int = 0
    start: int = 0
    end: int = 0
    section_path: str = ""


@dataclass(slots=True)
class ProcessedDocument:
    file_name: str
    source_format: str
    text_content: str
    markdown_content: str | None
    parser_strategy: str
    chunk_strategy: str
    chunks: list[ProcessedChunk]
    notes: list[str] = field(default_factory=list)


@dataclass(slots=True)
class _MineruConfig:
    api_token: str | None
    api_base: str
    lang: str
    enable_formula: bool
    enable_table: bool
    is_ocr: bool
    model_version: str | None
    request_timeout_s: float
    poll_attempts: int
    poll_interval_s: float

    @property
    def enabled(self) -> bool:
        return bool(self.api_token)


def normalize_parser_strategy(value: str | None) -> str:
    normalized = (value or PARSER_STRATEGY_AUTO).strip().lower()
    if normalized not in SUPPORTED_PARSER_STRATEGIES:
        return PARSER_STRATEGY_AUTO
    return normalized


def normalize_chunk_strategy(value: str | None) -> str:
    normalized = (value or CHUNK_STRATEGY_AUTO).strip().lower()
    if normalized not in SUPPORTED_CHUNK_STRATEGIES:
        return CHUNK_STRATEGY_AUTO
    return normalized


def processing_metadata_fields() -> dict[str, tuple[str, str]]:
    return {
        "parser_strategy": ("parser_strategy", PARSER_STRATEGY_AUTO),
        "chunk_strategy": ("chunk_strategy", CHUNK_STRATEGY_AUTO),
    }


async def process_knowledge_document(
    *,
    file_name: str,
    raw_bytes: bytes,
    chunk_size: int,
    chunk_overlap: int,
    separator: str,
    parser_strategy: str,
    chunk_strategy: str,
    mineru_overrides: dict[str, Any] | None = None,
) -> ProcessedDocument:
    parser_strategy = normalize_parser_strategy(parser_strategy)
    chunk_strategy = normalize_chunk_strategy(chunk_strategy)

    extension = Path(file_name).suffix.lower()
    notes: list[str] = []
    markdown_content: str | None = None
    text_content = ""

    if extension in {".md", ".markdown", ".mdx"}:
        markdown_content = raw_bytes.decode("utf-8", errors="replace").strip()
        text_content = markdown_content
    elif extension == ".docx" and parser_strategy != PARSER_STRATEGY_PLAIN_TEXT:
        markdown_content, text_content = await _docx_to_markdown_or_text(
            file_name=file_name,
            raw_bytes=raw_bytes,
            notes=notes,
            mineru_overrides=mineru_overrides,
        )
    elif extension == ".pdf" and parser_strategy != PARSER_STRATEGY_PLAIN_TEXT:
        markdown_content = await _pdf_to_markdown(file_name=file_name, raw_bytes=raw_bytes, notes=notes, mineru_overrides=mineru_overrides)
        if markdown_content:
            text_content = markdown_content

    if not text_content:
        text_content = _extract_text_with_fallbacks(file_name=file_name, raw_bytes=raw_bytes, notes=notes).strip()

    if not text_content and markdown_content:
        text_content = markdown_content.strip()

    effective_chunk_strategy = _resolve_chunk_strategy(
        requested=chunk_strategy,
        markdown_content=markdown_content,
    )

    if effective_chunk_strategy == CHUNK_STRATEGY_HEADING_MARKDOWN and markdown_content:
        chunks = _chunk_markdown_by_heading(
            markdown_content=markdown_content,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        if not chunks:
            effective_chunk_strategy = CHUNK_STRATEGY_RECURSIVE_TEXT
    else:
        chunks = []

    if effective_chunk_strategy == CHUNK_STRATEGY_RECURSIVE_TEXT:
        source_text = markdown_content or text_content
        chunks = _chunk_text_recursively(
            text=source_text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator=separator,
        )

    parser_used = parser_strategy
    if parser_used == PARSER_STRATEGY_AUTO:
        parser_used = PARSER_STRATEGY_MINERU_MARKDOWN if markdown_content else PARSER_STRATEGY_PLAIN_TEXT

    return ProcessedDocument(
        file_name=file_name,
        source_format=extension.lstrip(".") or "unknown",
        text_content=text_content,
        markdown_content=markdown_content,
        parser_strategy=parser_used,
        chunk_strategy=effective_chunk_strategy,
        chunks=chunks,
        notes=notes,
    )


def _resolve_chunk_strategy(*, requested: str, markdown_content: str | None) -> str:
    if requested == CHUNK_STRATEGY_HEADING_MARKDOWN:
        return CHUNK_STRATEGY_HEADING_MARKDOWN
    if requested == CHUNK_STRATEGY_RECURSIVE_TEXT:
        return CHUNK_STRATEGY_RECURSIVE_TEXT
    if markdown_content and _HEADING_RE.search(markdown_content):
        return CHUNK_STRATEGY_HEADING_MARKDOWN
    return CHUNK_STRATEGY_RECURSIVE_TEXT


async def _docx_to_markdown_or_text(*, file_name: str, raw_bytes: bytes, notes: list[str], mineru_overrides: dict[str, Any] | None = None) -> tuple[str | None, str]:
    mineru_markdown = await _docx_to_markdown_via_mineru(file_name=file_name, raw_bytes=raw_bytes, notes=notes, mineru_overrides=mineru_overrides)
    if mineru_markdown:
        return mineru_markdown, mineru_markdown

    direct_markdown = _docx_to_markdown_direct(raw_bytes)
    if direct_markdown:
        if not any("docx fallback" in note for note in notes):
            notes.append("docx fallback: used native DOCX to Markdown conversion.")
        return direct_markdown, direct_markdown

    text_content = _extract_text_with_fallbacks(file_name=file_name, raw_bytes=raw_bytes, notes=notes)
    return None, text_content


async def _docx_to_markdown_via_mineru(*, file_name: str, raw_bytes: bytes, notes: list[str], mineru_overrides: dict[str, Any] | None = None) -> str | None:
    config = _load_mineru_config(mineru_overrides)
    if not config.enabled:
        notes.append("docx fallback: Mineru is not configured.")
        return None

    with tempfile.TemporaryDirectory(prefix="earthmind-kb-docx-") as temp_dir:
        temp_root = Path(temp_dir)
        source_path = temp_root / file_name
        source_path.write_bytes(raw_bytes)

        pdf_path = _convert_office_to_pdf(source_path, temp_root, notes)
        if pdf_path is None:
            return None
        markdown, images, error = await _parse_pdf_with_mineru(pdf_path, config)
        if error:
            notes.append(f"docx fallback: Mineru parse failed ({error}).")
            return None
        if images:
            notes.append(f"Mineru extracted {len(images)} image(s) from {file_name}.")
        return _normalize_markdown(markdown or "")


async def _pdf_to_markdown(*, file_name: str, raw_bytes: bytes, notes: list[str], mineru_overrides: dict[str, Any] | None = None) -> str | None:
    config = _load_mineru_config(mineru_overrides)
    if not config.enabled:
        notes.append("pdf fallback: Mineru is not configured.")
        return None

    with tempfile.TemporaryDirectory(prefix="earthmind-kb-pdf-") as temp_dir:
        pdf_path = Path(temp_dir) / file_name
        pdf_path.write_bytes(raw_bytes)
        markdown, images, error = await _parse_pdf_with_mineru(pdf_path, config)
        if error:
            notes.append(f"pdf fallback: Mineru parse failed ({error}).")
            return None
        if images:
            notes.append(f"Mineru extracted {len(images)} image(s) from {file_name}.")
        return _normalize_markdown(markdown or "")


def _docx_to_markdown_direct(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        logger.debug("python-docx is unavailable for DOCX fallback: %s", exc)
        return ""

    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as exc:  # noqa: BLE001
        logger.debug("DOCX fallback open failed: %s", exc)
        return ""

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        style_name = ((paragraph.style.name if paragraph.style else "") or "").strip()
        match = re.search(r"heading\s+([1-6])", style_name, re.IGNORECASE)
        if match and text:
            lines.append(f"{'#' * int(match.group(1))} {text}")
            continue
        if text:
            lines.append(text)
        elif lines and lines[-1] != "":
            lines.append("")

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append(
                [
                    "\n".join((paragraph.text or "").strip() for paragraph in cell.paragraphs).strip()
                    for cell in row.cells
                ]
            )
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            continue
        if lines and lines[-1] != "":
            lines.append("")
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

    return _normalize_markdown("\n".join(lines))


def _extract_text_with_fallbacks(*, file_name: str, raw_bytes: bytes, notes: list[str]) -> str:
    try:
        extracted = extract_text_from_bytes(file_name, raw_bytes)
        if extracted and extracted.strip():
            return extracted
    except Exception as exc:  # noqa: BLE001
        notes.append(f"text extraction fallback: primary extractor failed ({exc}).")

    extension = Path(file_name).suffix.lower()
    if extension == ".pdf":
        return _extract_pdf_text_with_pypdf(raw_bytes)
    if extension == ".docx":
        return _extract_docx_text(raw_bytes)
    return raw_bytes.decode("utf-8", errors="replace")


def _extract_pdf_text_with_pypdf(raw_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # noqa: BLE001
        logger.debug("pypdf is unavailable for PDF fallback: %s", exc)
        return ""

    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
        parts = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        logger.debug("pypdf fallback failed: %s", exc)
        return ""


def _extract_docx_text(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        logger.debug("python-docx is unavailable for DOCX text fallback: %s", exc)
        return ""

    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as exc:  # noqa: BLE001
        logger.debug("DOCX text fallback open failed: %s", exc)
        return ""

    paragraphs = [(paragraph.text or "").strip() for paragraph in document.paragraphs]
    paragraphs = [text for text in paragraphs if text]
    return "\n\n".join(paragraphs).strip()


def _chunk_text_recursively(
    *,
    text: str,
    chunk_size: int,
    chunk_overlap: int,
    separator: str,
) -> list[ProcessedChunk]:
    if not text or not text.strip():
        return []

    separators = None
    if separator:
        actual_separator = separator.replace("\\n", "\n").replace("\\t", "\t")
        separators = [actual_separator, "\n\n", "\n", " ", ""]

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
    )
    chunks = splitter.split_text(text)
    processed: list[ProcessedChunk] = []
    position = 0
    for index, chunk in enumerate(chunks):
        start = text.find(chunk, position)
        if start < 0:
            start = position
        end = start + len(chunk)
        processed.append(
            ProcessedChunk(
                content=chunk,
                index=index,
                start=start,
                end=end,
            )
        )
        position = max(position, end - chunk_overlap)
    return processed


def _chunk_markdown_by_heading(
    *,
    markdown_content: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
) -> list[ProcessedChunk]:
    """Split markdown by heading sections, then sub-split oversized sections.

    Ported from AI_Writer's MarkdownChunker: each heading section is emitted
    as one chunk when it fits within *chunk_size*; sections that exceed the
    limit are further split by paragraph with sentence-boundary detection and
    overlap, mirroring AI_Writer's behaviour.
    """
    sections = _split_markdown_sections(markdown_content)
    if not sections:
        return []

    min_chunk_size = max(100, chunk_size // 10)
    allow_oversize = 0.2
    max_allowed = int(chunk_size * (1 + allow_oversize))

    processed: list[ProcessedChunk] = []
    cursor = 0
    for title, level, section_path, section_text in sections:
        prefix = f"{'#' * level} {title}\n\n" if level > 0 and title else ""
        if not section_text.strip() and not prefix:
            continue

        # Section fits in one chunk — emit directly.
        if len(section_text) <= chunk_size:
            content = f"{prefix}{section_text}".strip() if prefix else section_text.strip()
            if not content:
                content = prefix.strip() if prefix else title or ""
            if not content:
                continue
            start, end, cursor = _locate_chunk_bounds(markdown_content, content, cursor, 0)
            processed.append(
                ProcessedChunk(
                    content=content,
                    index=len(processed),
                    title=title,
                    level=level,
                    start=start,
                    end=end,
                    section_path=section_path,
                )
            )
            continue

        # Section too large — sub-split by paragraph with overlap + sentence
        # boundary detection (ported from AI_Writer MarkdownChunker._split_section).
        sub_chunks = _split_oversized_section(
            section_text=section_text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            min_chunk_size=min_chunk_size,
            max_allowed=max_allowed,
        )
        for sub_text in sub_chunks:
            content = f"{prefix}{sub_text}".strip() if prefix else sub_text.strip()
            if not content:
                continue
            start, end, cursor = _locate_chunk_bounds(markdown_content, content, cursor, chunk_overlap)
            processed.append(
                ProcessedChunk(
                    content=content,
                    index=len(processed),
                    title=title,
                    level=level,
                    start=start,
                    end=end,
                    section_path=section_path,
                )
            )

    return processed


def _split_oversized_section(
    *,
    section_text: str,
    chunk_size: int,
    chunk_overlap: int,
    min_chunk_size: int,
    max_allowed: int,
) -> list[str]:
    """Split a section whose body exceeds *chunk_size* into smaller pieces.

    Mirrors AI_Writer's MarkdownChunker._split_section: walks paragraphs,
    accumulates until the limit is reached, then flushes at a sentence
    boundary. Overlong individual paragraphs are force-split at sentence /
    whitespace boundaries with overlap.
    """
    paragraphs = re.split(r"\n\n+", section_text.strip())
    chunks: list[str] = []
    current_chunk = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # A single paragraph longer than the limit — force-split it.
        if len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""
            for piece in _split_long_paragraph(
                para,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                min_chunk_size=min_chunk_size,
            ):
                chunks.append(piece)
            continue

        potential_length = len(current_chunk) + len(para) + 2

        if potential_length <= chunk_size:
            current_chunk = para if not current_chunk else f"{current_chunk}\n\n{para}"
        elif potential_length <= max_allowed:
            # Allow slight overshoot to keep sentences intact.
            current_chunk = para if not current_chunk else f"{current_chunk}\n\n{para}"
            adjusted = _ensure_sentence_boundary(current_chunk, min_chunk_size)
            chunks.append(adjusted)
            remaining = current_chunk[len(adjusted):].strip()
            current_chunk = remaining if remaining else ""
        else:
            # Way over — flush current, start fresh with this paragraph.
            if current_chunk:
                adjusted = _ensure_sentence_boundary(current_chunk, min_chunk_size)
                chunks.append(adjusted)
                remaining = current_chunk[len(adjusted):].strip()
                current_chunk = f"{remaining}\n\n{para}" if remaining else para
            else:
                current_chunk = para

    if current_chunk:
        adjusted = _ensure_sentence_boundary(current_chunk, min_chunk_size)
        chunks.append(adjusted)
        remaining = current_chunk[len(adjusted):].strip()
        if remaining:
            chunks.append(remaining)

    return chunks


def _split_long_paragraph(
    paragraph: str,
    *,
    chunk_size: int,
    chunk_overlap: int,
    min_chunk_size: int,
) -> list[str]:
    """Force-split an overlong paragraph at sentence / whitespace boundaries."""
    chunks: list[str] = []
    start = 0
    while start < len(paragraph):
        end = start + chunk_size
        if end < len(paragraph):
            # Try sentence-ending punctuation first.
            for i in range(end, max(start + min_chunk_size, start), -1):
                if paragraph[i] in "。！？.!?":
                    end = i + 1
                    break
            else:
                # Fall back to whitespace.
                for i in range(end, max(start + min_chunk_size, start), -1):
                    if paragraph[i] in " \t":
                        end = i + 1
                        break
        piece = paragraph[start:end].strip()
        if piece:
            chunks.append(piece)
        start = end - chunk_overlap if end < len(paragraph) else end
    return chunks


def _ensure_sentence_boundary(text: str, min_chunk_size: int) -> str:
    """Truncate *text* so it ends at a complete sentence when possible."""
    sentence_endings = ("。", "！", "？", ".", "!", "?", "》", '"', "'", "）", ")", "】", "]")
    min_search = max(min_chunk_size, len(text) // 2)
    for i in range(len(text) - 1, min_search - 1, -1):
        if text[i] in sentence_endings:
            return text[: i + 1]
    secondary_marks = ("，", "、", ",", ";", "；", "：", ":", "（", "(", "【", "[")
    for i in range(len(text) - 1, min_search - 1, -1):
        if text[i] in secondary_marks:
            return text[: i + 1]
    return text


def _split_markdown_sections(markdown_content: str) -> list[tuple[str, int, str, str]]:
    lines = markdown_content.splitlines()
    sections: list[tuple[str, int, str, str]] = []
    stack: list[str] = []
    current_title = "Document Start"
    current_level = 0
    current_path = ""
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        text = "\n".join(current_lines).strip()
        if text or current_level > 0:
            sections.append((current_title, current_level, current_path, text))
        current_lines = []

    for line in lines:
        match = _HEADING_RE.match(line)
        if match:
            flush()
            level = len(match.group(1))
            title = match.group(2).strip()
            stack = stack[: level - 1]
            stack.append(title)
            current_title = title
            current_level = level
            current_path = " > ".join(stack)
            continue
        current_lines.append(line)

    flush()
    return sections


def _locate_chunk_bounds(
    full_text: str,
    snippet: str,
    cursor: int,
    chunk_overlap: int,
) -> tuple[int, int, int]:
    if not snippet:
        return cursor, cursor, cursor
    start = full_text.find(snippet, cursor)
    if start < 0:
        start = max(cursor, 0)
    end = start + len(snippet)
    next_cursor = max(end - chunk_overlap, 0)
    return start, end, next_cursor


def _normalize_markdown(markdown_content: str) -> str:
    normalized = markdown_content.replace("\r\n", "\n").replace("\r", "\n").strip()
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized


def _convert_office_to_pdf(source_path: Path, output_dir: Path, notes: list[str]) -> Path | None:
    soffice_cmd = _find_soffice()
    if soffice_cmd is None:
        notes.append("docx fallback: LibreOffice/soffice was not found for Mineru conversion.")
        return None

    try:
        result = subprocess.run(
            [
                soffice_cmd,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(source_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
    except Exception as exc:  # noqa: BLE001
        notes.append(f"docx fallback: LibreOffice conversion failed ({exc}).")
        return None

    if result.returncode != 0:
        stderr = (result.stderr or result.stdout or "").strip()
        notes.append(f"docx fallback: LibreOffice conversion failed ({stderr or 'unknown error'}).")
        return None

    pdf_path = output_dir / f"{source_path.stem}.pdf"
    if not pdf_path.exists():
        notes.append("docx fallback: converted PDF was not produced.")
        return None
    return pdf_path


def _find_soffice() -> str | None:
    for candidate in _SOFFICE_CANDIDATES:
        try:
            result = subprocess.run(
                [candidate, "--version"],
                capture_output=True,
                text=True,
                timeout=5,
                check=False,
            )
        except Exception:  # noqa: BLE001
            continue
        if result.returncode == 0:
            return candidate
    return None


def _env_flag(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() not in {"0", "false", "no", "off"}


def _load_mineru_config(overrides: dict[str, Any] | None = None) -> _MineruConfig:
    overrides = overrides or {}
    return _MineruConfig(
        api_token=overrides.get("mineru_api_token") or os.getenv("EARTHMIND_MINERU_API_TOKEN") or os.getenv("MINERU_API_TOKEN"),
        api_base=str(overrides.get("mineru_api_base") or os.getenv("EARTHMIND_MINERU_API_BASE", "https://mineru.net/api/v4")).rstrip("/"),
        lang=str(overrides.get("mineru_lang") or os.getenv("EARTHMIND_MINERU_LANG", "ch")),
        enable_formula=bool(overrides.get("mineru_enable_formula")) if overrides.get("mineru_enable_formula") is not None else _env_flag("EARTHMIND_MINERU_ENABLE_FORMULA", True),
        enable_table=bool(overrides.get("mineru_enable_table")) if overrides.get("mineru_enable_table") is not None else _env_flag("EARTHMIND_MINERU_ENABLE_TABLE", True),
        is_ocr=bool(overrides.get("mineru_is_ocr")) if overrides.get("mineru_is_ocr") is not None else _env_flag("EARTHMIND_MINERU_IS_OCR", True),
        model_version=str(overrides.get("mineru_model_version")) if overrides.get("mineru_model_version") else (os.getenv("EARTHMIND_MINERU_MODEL_VERSION") or None),
        request_timeout_s=float(overrides.get("mineru_timeout_s") or os.getenv("EARTHMIND_MINERU_TIMEOUT_S", "300")),
        poll_attempts=int(overrides.get("mineru_poll_attempts") or os.getenv("EARTHMIND_MINERU_POLL_ATTEMPTS", "180")),
        poll_interval_s=float(overrides.get("mineru_poll_interval_s") or os.getenv("EARTHMIND_MINERU_POLL_INTERVAL_S", "5")),
    )


async def _parse_pdf_with_mineru(
    pdf_path: Path,
    config: _MineruConfig,
) -> tuple[str | None, list[str], str | None]:
    if not config.enabled:
        return None, [], "Mineru API token is not configured."

    timeout = httpx.Timeout(config.request_timeout_s)
    async with httpx.AsyncClient(timeout=timeout) as client:
        upload_payload: dict[str, Any] = {
            "files": [
                {
                    "name": pdf_path.name,
                    "data_id": pdf_path.stem,
                    "is_ocr": config.is_ocr,
                }
            ],
            "enable_formula": config.enable_formula,
            "enable_table": config.enable_table,
            "language": config.lang,
        }
        if config.model_version:
            upload_payload["model_version"] = config.model_version

        batch_id, upload_url, error = await _request_mineru_upload_url(client, config, upload_payload)
        if error:
            return None, [], error

        error = await _upload_file_to_mineru(client, upload_url, pdf_path)
        if error:
            return None, [], error

        item, error = await _poll_mineru_batch(client, config, batch_id)
        if error:
            return None, [], error

        zip_url = (item or {}).get("full_zip_url")
        if not zip_url:
            return None, [], "Mineru response did not contain full_zip_url."

        return await _download_mineru_markdown(client, zip_url)


async def _request_mineru_upload_url(
    client: httpx.AsyncClient,
    config: _MineruConfig,
    payload: dict[str, Any],
) -> tuple[str | None, str | None, str | None]:
    headers = {
        "Authorization": f"Bearer {config.api_token}",
        "Content-Type": "application/json",
    }
    try:
        response = await client.post(f"{config.api_base}/file-urls/batch", json=payload, headers=headers)
        response.raise_for_status()
    except Exception as exc:  # noqa: BLE001
        return None, None, str(exc)

    body = response.json()
    if body.get("code") not in {0, 200}:
        return None, None, json.dumps(body, ensure_ascii=False)
    data = body.get("data") or {}
    urls = data.get("file_urls") or []
    batch_id = data.get("batch_id")
    if not batch_id or not urls:
        return None, None, "Mineru upload response was missing batch_id or file_urls."
    return str(batch_id), str(urls[0]), None


async def _upload_file_to_mineru(
    client: httpx.AsyncClient,
    upload_url: str,
    pdf_path: Path,
) -> str | None:
    try:
        response = await client.put(upload_url, content=pdf_path.read_bytes())
    except Exception as exc:  # noqa: BLE001
        return str(exc)
    if response.status_code != 200:
        return f"Mineru upload returned HTTP {response.status_code}."
    return None


async def _poll_mineru_batch(
    client: httpx.AsyncClient,
    config: _MineruConfig,
    batch_id: str,
) -> tuple[dict[str, Any] | None, str | None]:
    headers = {"Authorization": f"Bearer {config.api_token}"}
    url = f"{config.api_base}/extract-results/batch/{batch_id}"
    await asyncio.sleep(3)

    for _ in range(config.poll_attempts):
        await asyncio.sleep(config.poll_interval_s)
        try:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
        except Exception:  # noqa: BLE001
            continue
        body = response.json()
        results = ((body.get("data") or {}).get("extract_result")) or []
        if not results:
            continue
        item = results[0]
        state = item.get("state")
        if state == "done":
            return item, None
        if state == "failed":
            return None, str(item.get("err_msg") or "Mineru parsing failed.")
    return None, "Mineru parsing timed out."


async def _download_mineru_markdown(
    client: httpx.AsyncClient,
    zip_url: str,
) -> tuple[str | None, list[str], str | None]:
    try:
        response = await client.get(zip_url)
        response.raise_for_status()
    except Exception as exc:  # noqa: BLE001
        return None, [], str(exc)

    markdown_content: str | None = None
    images: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(response.content)) as archive:
            for name in archive.namelist():
                lower_name = name.lower()
                if markdown_content is None and lower_name.endswith("full.md"):
                    with archive.open(name) as file_obj:
                        markdown_content = file_obj.read().decode("utf-8", errors="ignore")
                elif markdown_content is None and lower_name.endswith(".md"):
                    with archive.open(name) as file_obj:
                        markdown_content = file_obj.read().decode("utf-8", errors="ignore")
                elif lower_name.startswith("images/") or "/images/" in lower_name:
                    images.append(Path(name).name)
    except Exception as exc:  # noqa: BLE001
        return None, [], str(exc)

    if not markdown_content:
        return None, [], "Mineru archive did not contain a markdown document."
    return _normalize_markdown(markdown_content), images, None
